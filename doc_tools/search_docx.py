import argparse
from docx import Document
#from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX 


def highlight_word_in_paragraphs(file_path, search_word):
    """
    This function opens a .docx file and highlights a word in a paragraph.
    NOTE: It does not highlight in tables, headers, footers, etc...
    
    :param file_path: Path to the .docx file
    :param search_word: Word to search and highlight
    """
    # Open the .docx file
    doc = Document(file_path)

    # Iterate over all paragraphs in the document
    for paragraph in doc.paragraphs:
        # Iterate over all runs in each paragraph
        for run in paragraph.runs:
            # If the search word is in the run text
            if search_word in run.text:
                # Change the font color to red to highlight it
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                
    doc.save('highlighted.docx')

def print_section_names(file_path):
    doc = Document(file_path)

    for paragraph in doc.paragraphs:
        if paragraph.style.name == 'Heading 1':
            print(paragraph.text)



if __name__ == "__main__":
    # Create the parser
    parser = argparse.ArgumentParser(description="Highlight a word in a .docx file")

    # Add the arguments
    parser.add_argument('File', metavar='file', type=str, help='the path to the .docx file')
    parser.add_argument('Word', metavar='word', type=str, help='the word to search and highlight')

    # Parse the arguments
    args = parser.parse_args()

    # Call the function with the arguments
    highlight_word_in_docx(args.File, args.Word)