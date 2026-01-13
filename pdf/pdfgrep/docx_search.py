#!usr/bin/python
# -*- coding: utf-8 -*-


from docx.text.paragraph import Paragraph
from docx.document import Document
from docx.table import _Cell, Table
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
import docx
import pandas as pd
import os

def get_table_names(document):
    table_names = []

    for paragraph in document.paragraphs:
        if paragraph.style.name.startswith('Caption') and paragraph.text.startswith('Table'):
            table_names.append(paragraph.text)
    
    return table_names


def iter_block_items(parent):
    if isinstance(parent, Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            table = Table(child, parent)
            for row in table.rows:
                for cell in row.cells:
                    yield from iter_block_items(cell)

def print_iter_blocks(document):
    for block in iter_block_items(document):
        print(block.text)


def process_highlights(document):

    highlights = []

    for paragraph in document.paragraphs:

        # Determine if it is a heading
        if paragraph.style.name.startswith('Heading'):
            last_heading_paragraph = paragraph

        highlight = ""
        
        for a_run in paragraph.runs:
            # Skip the headings that are highlighted
            if paragraph.style.name.startswith('Heading'):
                pass
            else:
                if a_run.font.highlight_color:
                    highlight += a_run.text
        
        if highlight:
            highlights.append((last_heading_paragraph.text, highlight))

    df = pd.DataFrame(highlights, columns = ['SECTION_NAME','SEARCH_TERM'])
    return df

def summarize_results (all_results_df):

    uniq_headings = list(all_results_df.SECTION_NAME.unique())

    results = []

    #Create a summary for each section's REGEX_MATCHES
    for a_heading in uniq_headings:
        # Get info about heading
        terms_df = df[df.SECTION_NAME == a_heading].SEARCH_TERM

        #unique search terms
        uniq_search_terms = list(terms_df.unique())

        uniq_search_terms_str = ','.join(uniq_search_terms)

        results.append((a_heading, uniq_search_terms_str))

    output_df = pd.DataFrame(results, columns=['SECTION_NAME', 'SEARCH_TERMS'])

    return output_df

if __name__ == "__main__":

    dirname = '/mnt/c/work/04_Projects/WPR00129344_SNMP_and_OM_CY25/docs/scratch/'
    filename = 'WID_SD_ICD_909_V11_HL_with_headings_and_tables.docx'
    document = docx.Document(os.path.join(dirname,filename))

    df = process_highlights(document)

    # Get unique headings 


    for a_heading in uniq_headings:
        




# for table in document.tables:
#     pass


# def iter_headings(paragraphs):
#     for paragraph in paragraphs:
#         if paragraph.style.name.startswith('Heading 1'):
#             yield paragraph

# for heading in iter_headings(document.paragraphs):
#     print(heading.text)




# for h in highlights:
#     print(h)

# document = opendocx(r'WID_SD_ICD_909_V11_HL.docx')
# words = document.xpath('//w:r', namespaces=document.nsmap)
# WPML_URI = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
# tag_rPr = WPML_URI + 'rPr'
# tag_highlight = WPML_URI + 'highlight'
# tag_val = WPML_URI + 'val'
# tag_t = WPML_URI + 't'
# for word in words:
#     for rPr in word.findall(tag_rPr):
#         high=rPr.findall(tag_highlight)
#         for hi in high:
#             if hi.attrib[tag_val] == 'yellow':
#                 print(word.find(tag_t).text.encode('utf-8').lower())